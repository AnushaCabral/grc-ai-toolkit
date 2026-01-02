"""
Document Export Utilities
"""

from typing import Dict, Any, Optional
from enum import Enum
from pathlib import Path
import logging
from datetime import datetime

try:
    from docx import Document as WordDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    """Supported export formats"""

    MARKDOWN = "markdown"
    WORD = "docx"
    PDF = "pdf"
    TEXT = "txt"
    HTML = "html"


class DocumentExporter:
    """
    Export documents to various formats

    Supports:
    - Markdown (.md)
    - Word (.docx) - requires python-docx
    - PDF (.pdf) - requires reportlab
    - Plain text (.txt)
    - HTML (.html)
    """

    @staticmethod
    def export_to_markdown(
        content: str,
        output_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Export to Markdown format

        Args:
            content: Document content
            output_path: Output file path
            metadata: Optional metadata to include

        Returns:
            Path to saved file
        """
        output_path = Path(output_path)

        # Add metadata as front matter if provided
        if metadata:
            front_matter = "---\n"
            for key, value in metadata.items():
                front_matter += f"{key}: {value}\n"
            front_matter += "---\n\n"
            content = front_matter + content

        # Write file
        output_path.write_text(content, encoding="utf-8")

        logger.info(f"Exported to Markdown: {output_path}")

        return output_path

    @staticmethod
    def export_to_word(
        content: str,
        output_path: str,
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Export to Word format

        Args:
            content: Document content (supports markdown-like formatting)
            output_path: Output file path
            title: Document title
            metadata: Optional metadata

        Returns:
            Path to saved file

        Raises:
            ImportError: If python-docx not installed
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        output_path = Path(output_path)

        # Create document
        doc = WordDocument()

        # Add title
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        if metadata:
            for key, value in metadata.items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))

            doc.add_paragraph()  # Blank line

        # Process content
        lines = content.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                doc.add_paragraph()
                continue

            # Headers
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)

            # Lists
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0:3].replace(".", "").isdigit():
                doc.add_paragraph(line.split(". ", 1)[1], style="List Number")

            # Bold/Italic (simple)
            elif "**" in line or "*" in line:
                p = doc.add_paragraph()
                parts = line.replace("**", "|").replace("*", "|").split("|")
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True

            # Normal paragraph
            else:
                doc.add_paragraph(line)

        # Save
        doc.save(str(output_path))

        logger.info(f"Exported to Word: {output_path}")

        return output_path

    @staticmethod
    def export_to_pdf(
        content: str,
        output_path: str,
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Export to PDF format

        Args:
            content: Document content
            output_path: Output file path
            title: Document title
            metadata: Optional metadata

        Returns:
            Path to saved file

        Raises:
            ImportError: If reportlab not installed
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError(
                "reportlab not installed. Install with: pip install reportlab"
            )

        output_path = Path(output_path)

        # Create PDF
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Add title
        if title:
            story.append(Paragraph(title, styles["Title"]))
            story.append(Spacer(1, 12))

        # Add metadata
        if metadata:
            for key, value in metadata.items():
                text = f"<b>{key}:</b> {value}"
                story.append(Paragraph(text, styles["Normal"]))
            story.append(Spacer(1, 12))

        # Process content
        lines = content.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                story.append(Spacer(1, 6))
                continue

            # Headers
            if line.startswith("# "):
                story.append(Paragraph(line[2:], styles["Heading1"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], styles["Heading3"]))

            # Normal text
            else:
                # Escape HTML entities and handle basic markdown
                line = line.replace("<", "&lt;").replace(">", "&gt;")
                line = line.replace("**", "<b>").replace("**", "</b>")
                line = line.replace("*", "<i>").replace("*", "</i>")

                story.append(Paragraph(line, styles["Normal"]))

        # Build PDF
        doc.build(story)

        logger.info(f"Exported to PDF: {output_path}")

        return output_path

    @staticmethod
    def export_to_text(
        content: str,
        output_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Export to plain text format

        Args:
            content: Document content
            output_path: Output file path
            metadata: Optional metadata

        Returns:
            Path to saved file
        """
        output_path = Path(output_path)

        # Add metadata header
        if metadata:
            header = "\n".join(f"{key}: {value}" for key, value in metadata.items())
            content = f"{header}\n\n{'=' * 80}\n\n{content}"

        # Write file
        output_path.write_text(content, encoding="utf-8")

        logger.info(f"Exported to text: {output_path}")

        return output_path

    @staticmethod
    def export_to_html(
        content: str,
        output_path: str,
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Export to HTML format

        Args:
            content: Document content (markdown-like)
            output_path: Output file path
            title: Document title
            metadata: Optional metadata

        Returns:
            Path to saved file
        """
        output_path = Path(output_path)

        # Build HTML
        html = ["<!DOCTYPE html>", "<html>", "<head>", "<meta charset='UTF-8'>"]

        if title:
            html.append(f"<title>{title}</title>")

        html.append(
            """
<style>
    body { font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }
    h1 { color: #2c3e50; border-bottom: 2px solid #3498db; }
    h2 { color: #34495e; margin-top: 30px; }
    h3 { color: #7f8c8d; }
    code { background: #ecf0f1; padding: 2px 6px; border-radius: 3px; }
    pre { background: #ecf0f1; padding: 15px; border-radius: 5px; overflow-x: auto; }
    .metadata { background: #f8f9fa; padding: 10px; border-left: 3px solid #3498db; margin-bottom: 20px; }
</style>
"""
        )

        html.append("</head>", "<body>")

        # Add title
        if title:
            html.append(f"<h1>{title}</h1>")

        # Add metadata
        if metadata:
            html.append("<div class='metadata'>")
            for key, value in metadata.items():
                html.append(f"<p><strong>{key}:</strong> {value}</p>")
            html.append("</div>")

        # Process content
        lines = content.split("\n")

        in_list = False

        for line in lines:
            line = line.strip()

            if not line:
                if in_list:
                    html.append("</ul>")
                    in_list = False
                html.append("<br>")
                continue

            # Headers
            if line.startswith("# "):
                html.append(f"<h1>{line[2:]}</h1>")
            elif line.startswith("## "):
                html.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith("### "):
                html.append(f"<h3>{line[4:]}</h3>")

            # Lists
            elif line.startswith("- ") or line.startswith("* "):
                if not in_list:
                    html.append("<ul>")
                    in_list = True
                html.append(f"<li>{line[2:]}</li>")

            # Normal paragraph
            else:
                if in_list:
                    html.append("</ul>")
                    in_list = False

                # Handle basic markdown
                line = line.replace("**", "<strong>").replace("**", "</strong>")
                line = line.replace("*", "<em>").replace("*", "</em>")
                line = line.replace("`", "<code>").replace("`", "</code>")

                html.append(f"<p>{line}</p>")

        if in_list:
            html.append("</ul>")

        html.append("</body>", "</html>")

        # Write file
        output_path.write_text("\n".join(html), encoding="utf-8")

        logger.info(f"Exported to HTML: {output_path}")

        return output_path

    @classmethod
    def export(
        cls,
        content: str,
        output_path: str,
        format: ExportFormat,
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Export document to specified format

        Args:
            content: Document content
            output_path: Output file path
            format: Export format
            title: Document title
            metadata: Optional metadata

        Returns:
            Path to saved file
        """
        # Add default metadata
        if metadata is None:
            metadata = {}

        metadata.setdefault("generated_date", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        metadata.setdefault("generator", "GRC AI Toolkit")

        # Export based on format
        if format == ExportFormat.MARKDOWN:
            return cls.export_to_markdown(content, output_path, metadata)

        elif format == ExportFormat.WORD:
            return cls.export_to_word(content, output_path, title, metadata)

        elif format == ExportFormat.PDF:
            return cls.export_to_pdf(content, output_path, title, metadata)

        elif format == ExportFormat.TEXT:
            return cls.export_to_text(content, output_path, metadata)

        elif format == ExportFormat.HTML:
            return cls.export_to_html(content, output_path, title, metadata)

        else:
            raise ValueError(f"Unsupported export format: {format}")
